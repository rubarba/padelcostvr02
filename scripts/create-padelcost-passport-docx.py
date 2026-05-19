from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "/Users/ruibarba/Desktop/files 2/PadelCost-passaporte-tecnico-extremamente-detalhado.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(18, 31, 49)
MUTED = RGBColor(91, 111, 135)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color


def set_table_borders(table, color="D8E0EA", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_column_widths(table, widths_cm):
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else DARK_BLUE
    return p


def add_para(doc, text="", style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        rest = text[len(bold_lead):]
        if rest:
            p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        p.add_run(item)


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [16.5])
    set_table_borders(table, color="D8E0EA", size="6")
    set_cell_margins(table, top=140, bottom=140, start=180, end=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = INK
    doc.add_paragraph()


def add_kv_table(doc, rows, widths=(4.2, 12.3), header=None):
    extra = 1 if header else 0
    table = doc.add_table(rows=len(rows) + extra, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, list(widths))
    set_table_borders(table)
    set_cell_margins(table)
    start = 0
    if header:
        table.cell(0, 0).merge(table.cell(0, 1))
        set_cell_shading(table.cell(0, 0), LIGHT_BLUE)
        set_cell_text(table.cell(0, 0), header, bold=True, color=DARK_BLUE)
        start = 1
    for i, (k, v) in enumerate(rows, start=start):
        c0, c1 = table.cell(i, 0), table.cell(i, 1)
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c0, LIGHT_GRAY)
        set_cell_text(c0, k, bold=True, color=DARK_BLUE)
        set_cell_text(c1, v)
    doc.add_paragraph()
    return table


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        set_column_widths(table, widths)
    set_table_borders(table)
    set_cell_margins(table)
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_text(cell, h, bold=True, color=DARK_BLUE)
    for r_i, row in enumerate(rows, start=1):
        for c_i, value in enumerate(row):
            cell = table.cell(r_i, c_i)
            set_cell_text(cell, value)
    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [16.5])
    set_table_borders(table, color="CFD8E3", size="4")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        s = styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        s = styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(11)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.25
        s.paragraph_format.left_indent = Cm(0.95)
        s.paragraph_format.first_line_indent = Cm(-0.48)


doc = Document()
configure_styles(doc)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.add_run("PadelCost — Passaporte Técnico Extremamente Detalhado")
subtitle = doc.add_paragraph(style="Subtitle")
subtitle.add_run("Documentação interna para continuidade do projeto em nova conversa ou por outro programador")

add_kv_table(doc, [
    ("Projeto", "PadelCost.pt"),
    ("Pasta local de trabalho", "/Users/ruibarba/Desktop/files 2"),
    ("Publicação", "Manual para GitHub, repositório PadelCost VR2"),
    ("Hosting", "GitHub Pages, site estático sem backend"),
    ("Estado do documento", "Gerado em 2026-05-19, após integração da Deporvillage"),
    ("Nota de segurança", "Não colocar scripts/.env, API keys, URLs privados completos ou ficheiros temporários no GitHub."),
], header="Metadados operacionais")

add_callout(
    doc,
    "Resumo executivo",
    "PadelCost é um comparador estático de preços de produtos de padel. A aplicação principal vive em index.html, lê ficheiros JS gerados em data/, agrega ofertas por produto e permite comparar preços entre lojas parceiras. Os feeds são processados localmente por scripts Node.js e só os ficheiros públicos gerados são publicados no GitHub Pages."
)

add_heading(doc, "Índice de Conteúdos", 1)
add_numbered(doc, [
    "Visão geral do projeto",
    "Arquitetura técnica",
    "Lógica do sistema",
    "UX/UI e design system",
    "Componentes principais",
    "Estado atual do desenvolvimento",
    "Problemas e limitações atuais",
    "Próximos passos",
    "Regras importantes",
    "Linguagem e padrões de código",
    "Runbook operacional",
    "Resumo essencial",
])

add_heading(doc, "1) Visão Geral do Projeto", 1)
add_para(doc, "Objetivo do site: criar um comparador de preços de artigos de padel para o mercado português/ibérico, permitindo ao utilizador encontrar rapidamente em que loja um produto está mais barato e perceber se existe comparação real entre várias lojas.")
add_para(doc, "Público-alvo: jogadores de padel recreativos e avançados, compradores de raquetes, sapatilhas, sacos e acessórios, utilizadores portugueses que pesquisam preços antes de comprar, e potenciais leitores vindos do Google através de páginas SEO de categorias/produtos.")
add_para(doc, "Proposta de valor: o PadelCost não é uma loja; é uma camada de comparação e descoberta. Junta produtos de várias lojas parceiras, normaliza nomes/categorias, mostra preços e links de saída para a loja, e ajuda a reduzir a fricção de pesquisar manualmente em várias lojas.")
add_para(doc, "Estado atual: MVP funcional em desenvolvimento ativo. O site já tem home/catalogo, filtros, modal de produto, favoritos locais, páginas SEO estáticas, páginas de categoria, páginas de produto geradas, página de lojas, política/cookies/contacto/guias, pipeline de feeds, relatórios de qualidade e ferramenta manual de revisão de duplicados.")

add_kv_table(doc, [
    ("Modelo de negócio", "Afiliados. Links de loja podem ser links afiliados provenientes de AWIN, TradeTracker ou feeds/parâmetros próprios."),
    ("Categorias públicas", "Raquetes, sapatilhas, sacos de padel e acessórios."),
    ("Categorias brutas internas", "Incluem também roupa, mas roupa não é foco público atual e pode ficar fora da UI principal."),
    ("Número atual de lojas integradas", "8: Padel Market, Zona de Padel, Padel Proshop PT, Atmosfera Sport, Forum Sport ES, Decathlon ES, Adidas Padel e Deporvillage."),
    ("Publicação", "O utilizador gera localmente, valida e depois envia manualmente os ficheiros para GitHub."),
], header="Estado funcional")

add_heading(doc, "2) Arquitetura Técnica", 1)
add_para(doc, "A base de desenvolvimento está na pasta local /Users/ruibarba/Desktop/files 2. Esta pasta contém o site estático, os dados gerados, os scripts Node.js para feeds e os assets visuais. O GitHub Pages serve os ficheiros estáticos; não existe backend, base de dados remota, API server, autenticação nem jobs automáticos em produção.")
add_para(doc, "Tecnologias usadas: HTML, CSS e JavaScript em ficheiro estático; React/Babel carregado no frontend dentro do index.html; Font Awesome para ícones; ficheiros JS de dados que populam variáveis globais window.*; Node.js para scripts locais de geração; csv-parse/dotenv para ingestão de feeds; GitHub Pages para hosting.")

add_matrix(doc, ["Área", "Ficheiros/pastas", "Função"], [
    ("Aplicação principal", "index.html", "Catálogo, home, filtros, modais, favoritos, comparação e UI principal."),
    ("Páginas institucionais", "como-funciona.html, contacto.html, cookies.html, privacidade.html, dados.html", "Conteúdo estático de suporte, confiança e conformidade."),
    ("Páginas SEO", "categoria/*.html e produto/*.html", "Páginas indexáveis geradas automaticamente por scripts/generate-seo-pages.js."),
    ("Dados públicos", "data/*.js", "Ficheiros JavaScript com arrays de produtos/ofertas expostos ao frontend."),
    ("Scripts locais", "scripts/*.js", "Ingestão de feeds, normalização, merge, relatórios, SEO e revisão de duplicados."),
    ("Assets visuais", "images/, logos/, logos_marcas/", "Banners, cards, logotipos de lojas e marcas."),
    ("Ferramentas de revisão", "duplicados-revisao.html, duplicados-revisao.csv, data/duplicate-decisions.json", "Gestão manual de possíveis duplicados."),
], widths=[3.3, 5.3, 7.9])

add_heading(doc, "2.1 Estrutura de ficheiros atual", 2)
add_code_block(doc, [
    "/Users/ruibarba/Desktop/files 2/",
    "  index.html",
    "  lojas.html",
    "  categoria/",
    "    raquetes-de-padel.html",
    "    sapatilhas-de-padel.html",
    "    sacos-de-padel.html",
    "    acessorios-de-padel.html",
    "  produto/",
    "    *.html  (páginas de produto geradas)",
    "  data/",
    "    products-data.js",
    "    seo-pages-data.js",
    "    duplicate-decisions.json",
    "    adidas-padel-data.js",
    "    padel-market-data.js",
    "    padel-proshop-data.js",
    "    forum-sport-data.js",
    "    zona-de-padel-data.js",
    "    decathlon-data.js",
    "    deporvillage-data.js",
    "    amazon-manual-data.js  (histórico/manual, não é fonte prioritária atual)",
    "  scripts/",
    "    generate-*.js",
    "    merge-offers.js",
    "    report-quality.js",
    "    report-duplicate-candidates.js",
    "    generate-duplicate-review-page.js",
    "    generate-seo-pages.js",
    "    update-weekly-no-atmosfera.js",
    "    .env  (privado, nunca publicar)",
    "    .env.example",
    "  images/",
    "  logos/",
    "  logos_marcas/",
    "  sitemap.xml",
    "  404.html",
])

add_heading(doc, "2.2 Como os dados são carregados", 2)
add_para(doc, "O frontend não chama feeds externos em tempo real. O index.html carrega ficheiros JS locais/publicados, principalmente data/products-data.js e data/seo-pages-data.js. Estes ficheiros definem variáveis globais como window.PADELCOST_PRODUCTS e window.PADELCOST_SEO_PRODUCT_URLS. Depois, o React embutido no index.html lê essas variáveis e constrói o catálogo.")
add_para(doc, "Os feeds privados são descarregados e processados localmente pelos scripts em /scripts. As chaves e URLs privados ficam em scripts/.env. Os dados resultantes são limpos e escritos em /data/*.js. Só estes ficheiros gerados, sem segredos, vão para GitHub Pages.")

add_heading(doc, "2.3 Dependências externas e lojas", 2)
add_matrix(doc, ["Loja", "Fonte atual", "Ficheiro gerado", "Notas"], [
    ("Padel Market", "AWIN", "data/padel-market-data.js", "Feed com raquetes, sapatilhas e sacos. Categorias foram revistas para não cortar produtos legítimos."),
    ("Padel Proshop PT", "Feed externo/loja parceira", "data/padel-proshop-data.js", "Loja com muitos produtos e boa cobertura; usada fortemente no merge."),
    ("Zona de Padel", "Google Merchant XML / feed", "data/zona-de-padel-data.js", "Boa cobertura de raquetes/sacos/sapatilhas/acessórios; preços precisam de atualização periódica."),
    ("Forum Sport ES", "Feed externo/AWIN", "data/forum-sport-data.js", "Foi necessário afinar regras para não cortar produtos legítimos."),
    ("Atmosfera Sport", "Feed externo/AWIN", "data/products-data.js como base histórica + ofertas", "Feed com problemas conhecidos; rotina semanal atual evita regenerar Atmosfera até revisão."),
    ("Adidas Padel", "TradeTracker/feed próprio", "data/adidas-padel-data.js", "Loja oficial Adidas; preços/stock podem diferir de retalhistas."),
    ("Decathlon ES", "AWIN", "data/decathlon-data.js", "Integrada, mas filtro mantém só produtos vendidos pela Decathlon e relevantes."),
    ("Deporvillage", "AWIN", "data/deporvillage-data.js", "Integrada em 2026-05-19. Feed tem 68 raquetes, 6 sacos, 5 acessórios após filtros seguros."),
], widths=[3.0, 3.4, 4.7, 5.4])

add_callout(doc, "Segurança de feeds", "Nunca colocar URLs completos com API key dentro do frontend, README público, GitHub ou documentação partilhada. Usar variáveis no scripts/.env e publicar apenas ficheiros data/*.js já processados.")

add_heading(doc, "3) Lógica do Sistema", 1)
add_heading(doc, "3.1 Modelo mental do comparador", 2)
add_para(doc, "O sistema separa produto comparável de oferta de loja. Um produto é o item conceptual mostrado ao utilizador: por exemplo, uma raquete Adidas Metalbone Team Light 2026. Uma oferta é a presença desse produto numa loja específica com preço, stock, URL e possível custo de entrega. O objetivo do merge é juntar ofertas equivalentes no mesmo produto para mostrar “Ver 2 lojas”, “Ver 5 lojas”, etc.")
add_para(doc, "Quando o produto só tem uma oferta, o site mostra o nome da loja no botão do card e uma mensagem no modal: “Oferta única neste momento. Estamos a adicionar mais lojas para comparação.” Isto reduz frustração e explica que o comparador ainda está a expandir cobertura.")

add_heading(doc, "3.2 Estrutura de dados principal", 2)
add_para(doc, "O ficheiro data/products-data.js contém um array de produtos em window.PADELCOST_PRODUCTS. A estrutura real pode variar por loja/categoria, mas o contrato prático é:")
add_matrix(doc, ["Campo", "Tipo", "Função"], [
    ("id", "string/number", "Identificador interno do produto agregado."),
    ("name", "string", "Nome normalizado e exibido ao utilizador."),
    ("brand", "string", "Marca normalizada, usada em filtros e cards."),
    ("category", "string", "raquetes, sapatilhas, sacos, acessorios ou roupa."),
    ("price", "number", "Melhor preço atual do produto após ordenar as lojas."),
    ("oldPrice", "number|null", "Preço anterior/PVP quando disponível."),
    ("image", "string", "URL ou caminho da imagem principal."),
    ("imageSource/source", "string", "Origem do produto/imagem, útil para scoring e diagnóstico."),
    ("ean/productGTIN/mpn/upc", "string|null", "Identificadores usados para merge seguro."),
    ("sourceProductId", "string|null", "ID do produto na loja/feed de origem."),
    ("sourceCategory", "string|null", "Categoria original do feed."),
    ("description", "string|null", "Descrição original/normalizada para modal e SEO."),
    ("specs", "object", "Características extra: peso, forma, nível, material, tamanhos, tipo, etc."),
    ("stores", "array", "Lista de ofertas por loja: nome, preço, stock, URL, custos, etc."),
], widths=[3.4, 3.0, 10.1])

add_heading(doc, "3.3 Estrutura de uma oferta em stores[]", 2)
add_matrix(doc, ["Campo", "Função"], [
    ("key", "Chave técnica da loja, por exemplo deporvillage ou padel-market."),
    ("name", "Nome visível: Padel Market, Zona de Padel, Deporvillage, etc."),
    ("price", "Preço numérico usado para ordenar e escolher melhor oferta."),
    ("stock", "Texto de stock normalizado quando o feed permite: Em stock, Sem stock ou Disponibilidade por confirmar."),
    ("url", "Link de saída para a loja. Pode ser link afiliado."),
    ("deliveryCost", "Custo de entrega quando vem no feed; nem sempre existe."),
])

add_heading(doc, "3.4 Pipeline de geração e merge", 2)
add_numbered(doc, [
    "Cada loja tem um script gerador próprio em scripts/generate-*.js.",
    "O script descarrega/lê o feed, parseia CSV/XML conforme a loja, normaliza campos, mapeia categorias e aplica filtros anti-intrusos.",
    "O resultado por loja é escrito em data/<loja>-data.js com uma variável global própria.",
    "scripts/merge-offers.js lê data/products-data.js e os ficheiros por loja.",
    "Antes de reintegrar lojas geradas, o merge remove ofertas antigas dessas lojas para evitar preços desatualizados.",
    "O merge tenta casar ofertas por EAN, GTIN, MPN e, em casos seguros, assinatura de nome/categoria.",
    "Se encontrar produto equivalente, adiciona/atualiza a loja no stores[]. Se não encontrar, cria novo produto.",
    "Aplica decisões manuais de duplicados guardadas em data/duplicate-decisions.json.",
    "Remove produtos sem lojas e escreve data/products-data.js atualizado.",
    "Depois correm report-quality.js, generate-seo-pages.js e report-duplicate-candidates.js.",
])

add_heading(doc, "3.5 Regras de matching", 2)
add_bullets(doc, [
    ("EAN/GTIN primeiro: ", "quando dois feeds trazem identificadores fortes, estes são a forma mais segura de juntar produtos."),
    ("MPN depois: ", "útil quando EAN falha, mas pode variar por país/cor/tamanho; por isso é usado com cuidado."),
    ("Assinatura de nome: ", "fallback só quando a assinatura por categoria/marca/nome é suficientemente segura e única."),
    ("Decisões manuais: ", "data/duplicate-decisions.json guarda pares aprovados ou rejeitados para não perder histórico quando o feed atualiza."),
    ("Evitar regras demasiado apertadas: ", "a preferência atual é não cortar produtos legítimos; quando houver dúvida, gerar candidatos de revisão manual."),
])

add_heading(doc, "3.6 Contagens atuais", 2)
add_matrix(doc, ["Métrica", "Valor atual"], [
    ("Produtos totais em data/products-data.js", "2544"),
    ("Produtos públicos", "2353"),
    ("Produtos core visíveis/agregados", "1560"),
    ("Produtos fora do público", "191, sobretudo roupa"),
    ("Produtos com 2+ lojas nas categorias públicas", "461"),
    ("Produtos com apenas 1 loja nas categorias públicas", "1892"),
    ("Páginas de produto SEO geradas", "442"),
    ("URLs no sitemap", "458"),
    ("Candidatos de duplicado por decidir após Deporvillage", "12"),
])

add_matrix(doc, ["Loja", "Ofertas no catálogo agregado", "Categorias agregadas"], [
    ("Padel Market", "747", "102 raquetes, 596 sapatilhas, 49 sacos"),
    ("Zona de Padel", "592", "226 raquetes, 185 sapatilhas, 128 sacos, 53 acessórios"),
    ("Padel Proshop PT", "571", "229 raquetes, 102 sapatilhas, 151 sacos, 89 acessórios"),
    ("Atmosfera Sport", "663", "107 raquetes, 190 sapatilhas, 97 sacos, 78 acessórios, 191 roupa"),
    ("Forum Sport ES", "324", "139 raquetes, 150 sapatilhas, 18 sacos, 17 acessórios"),
    ("Decathlon ES", "213", "79 raquetes, 77 sapatilhas, 11 sacos, 46 acessórios"),
    ("Adidas Padel", "139", "79 raquetes, 11 sapatilhas, 49 sacos"),
    ("Deporvillage", "79", "68 raquetes, 6 sacos, 5 acessórios"),
], widths=[3.5, 3.6, 9.4])

add_heading(doc, "4) UX/UI e Design System", 1)
add_para(doc, "O ADN visual do PadelCost combina utilidade de comparador com linguagem premium/desportiva. A marca usa fundo escuro/azul, contraste forte, botões azuis arredondados, cartões brancos sobre textura clara e assets com estética técnica/cinematográfica.")
add_bullets(doc, [
    ("Estilo geral: ", "dark navy no topo, azul vivo para CTAs, branco/cinzento claro para área de catálogo, tipografia forte e legível."),
    ("ADN das imagens: ", "noir, blueprint técnico, ambiente de padel premium, imagens de produto claras e inspecionáveis."),
    ("Sem marcas em assets editoriais: ", "as imagens decorativas/guias devem evitar marcas comerciais visíveis quando não forem produtos reais do feed."),
    ("Cards de produto: ", "imagem no topo, coração/favorito, marca em caixa alta, nome, preço, botão principal, link de comparar."),
    ("Botões por número de lojas: ", "produtos com 2+ lojas usam azul principal; produtos com uma loja usam azul mais escuro e o nome da loja no botão."),
    ("Mensagem de oferta única: ", "no modal aparece aviso para reduzir frustração quando só existe uma loja."),
    ("Cards editoriais/guias: ", "há cartões de Perfil Jogador, Raquetes, Sapatilhas e guias com imagem forte e copy curta."),
])

add_heading(doc, "4.1 Categorias e navegação", 2)
add_para(doc, "A navegação principal por categoria tem Raquetes, Sapatilhas, Sacos de padel e Acessórios. Estas categorias existem na home/app e também como páginas SEO estáticas em /categoria/. As categorias na home são mais interativas, com filtros e paginação; as páginas SEO são mais simples, indexáveis e orientadas a pesquisa orgânica.")
add_para(doc, "Existem dois “tipos” de página de categoria: a categoria dinâmica dentro de index.html e as páginas SEO estáticas em categoria/*.html. A dinâmica é para utilizador explorar; a estática é para Google indexar e para captar pesquisas de categoria.")

add_heading(doc, "5) Componentes Principais", 1)
add_matrix(doc, ["Componente", "Onde vive", "Descrição técnica"], [
    ("Header/topbar", "index.html", "Logo PadelCost, claim, pesquisa, favoritos e layout responsivo."),
    ("Tabs de categoria", "index.html", "Controlam activeCategory e limpam filtros/paginação ao mudar."),
    ("Filtros", "index.html", "Marca, loja, preço máximo e filtros específicos por categoria: forma/nível, género/tamanho, tipo."),
    ("Ordenação", "index.html", "Relevância, preço ascendente/descendente, rating e desconto. Relevância privilegia produtos com 2+ lojas."),
    ("Cards de produto", "index.html", "Mostram produto, melhor preço, botão da loja ou número de lojas, comparar e favoritos."),
    ("Modal de produto", "index.html", "Imagem, marca, título, descrição, specs compactas, lojas ordenadas por preço, botões para loja."),
    ("Comparador", "index.html", "Permite selecionar produtos para comparação lado a lado."),
    ("Favoritos", "index.html + localStorage", "Guardados apenas localmente quando cookies são aceites."),
    ("Página Lojas", "lojas.html", "Agrega produtos por loja e mostra lojas ativas/produtos monitorizados."),
    ("Páginas SEO", "scripts/generate-seo-pages.js", "Gera páginas produto/categoria e sitemap.xml."),
    ("Revisão de duplicados", "duplicados-revisao.html", "UI local para aprovar/rejeitar pares e exportar decisões JSON."),
], widths=[3.2, 4.3, 9.0])

add_heading(doc, "6) Estado Atual do Desenvolvimento", 1)
add_heading(doc, "6.1 Já está feito", 2)
add_bullets(doc, [
    "Catálogo principal funcional em index.html.",
    "Categorias públicas: raquetes, sapatilhas, sacos e acessórios.",
    "Filtros por marca, loja, preço e características específicas.",
    "Ordenação inteligente que dá prioridade a produtos com 2+ lojas.",
    "Botões diferenciados para produto com uma loja versus várias lojas.",
    "Modal com lista de lojas em formato compacto e mensagem de oferta única.",
    "Página de lojas com contagem agregada e logos.",
    "Páginas SEO de categorias e produtos comparáveis.",
    "Sitemap atualizado automaticamente pelo gerador SEO.",
    "Scripts por loja para Adidas, Padel Market, Padel Proshop, Forum Sport ES, Zona de Padel, Decathlon ES e Deporvillage.",
    "Relatório de qualidade sem problemas bloqueantes atuais.",
    "Ferramenta local de duplicados com histórico em JSON.",
    "Integração Deporvillage concluída com 79 ofertas úteis.",
])

add_heading(doc, "6.2 Parcialmente feito ou em atenção", 2)
add_bullets(doc, [
    "Atmosfera Sport tem histórico de problemas no feed e não deve ser regenerada na rotina semanal até revisão.",
    "Amazon foi removida/abandonada como fonte regular porque eram links manuais e ficavam desatualizados.",
    "As páginas SEO existem, mas o trabalho SEO ainda deve ser aprofundado: títulos, meta descriptions, dados estruturados, conteúdo de apoio e interlinking.",
    "A revisão de duplicados é manual; o sistema deteta candidatos, mas o utilizador tem de decidir pares ambíguos.",
    "As contagens de loja e categoria podem diferir entre feed bruto, catálogo agregado e UI porque há deduplicação e filtros de intrusos.",
])

add_heading(doc, "7) Problemas e Limitações Atuais", 1)
add_bullets(doc, [
    ("GitHub Pages sem backend: ", "não há atualização automática em produção; o utilizador tem de correr scripts localmente e publicar ficheiros gerados."),
    ("Feeds externos instáveis: ", "lojas podem mudar nomes, preços, colunas, categorias ou stock sem aviso."),
    ("Diferença site vs feed: ", "o site de uma loja pode mostrar mais produtos do que o feed afiliado disponibiliza; o PadelCost só deve usar o feed permitido."),
    ("Matching imperfeito: ", "EAN/GTIN/MPN podem divergir por país, variante, cor ou marketplace. Por isso há revisão manual."),
    ("Atmosfera Sport: ", "não regenerar automaticamente até resolver inconsistências; manter com cautela."),
    ("Acessórios: ", "categoria sensível porque bolas, grips, protetores, sacos e roupa podem vir misturados nos feeds."),
    ("Decathlon/Worten/marketplaces: ", "em lojas com marketplace, é preciso decidir se só entram produtos vendidos pela própria loja. Para Worten foi concluído que não valia a pena integrar naquele momento."),
    ("Dados estruturados Google: ", "Search Console assinalou aggregateRating/review como melhoria não crítica; não é bloqueante. Não inventar reviews/rating falsos."),
])

add_heading(doc, "8) Próximos Passos", 1)
add_heading(doc, "8.1 Prioridades técnicas imediatas", 2)
add_numbered(doc, [
    "Rever os 12 duplicados atuais em duplicados-revisao.html e importar/exportar decisões para data/duplicate-decisions.json.",
    "Depois de cada decisão de duplicados, correr merge, relatório de qualidade, SEO e revisão de duplicados novamente.",
    "Continuar integração de novas lojas apenas com feed permitido: Sport is Good será a próxima loja quando o feed for fornecido.",
    "Aprofundar SEO: títulos, meta descriptions, dados estruturados válidos, conteúdo de categoria e interlinking.",
    "Criar ou melhorar documentação operacional com lista exata de ficheiros a publicar no GitHub.",
])

add_heading(doc, "8.2 Melhorias UX", 2)
add_bullets(doc, [
    "Melhorar estados vazios quando só existe uma loja ou quando filtros reduzem resultados.",
    "Refinar mobile da página de produto/modal para listas de lojas mais compactas.",
    "Adicionar indicação visual clara de “comparável” versus “oferta única”.",
    "Melhorar leitura das características em mobile, mantendo-as compactas.",
    "Avaliar paginação e performance se o catálogo continuar a crescer.",
])

add_heading(doc, "8.3 Integrações futuras", 2)
add_bullets(doc, [
    "Sport is Good, quando o feed for enviado.",
    "Outras lojas AWIN/TradeTracker com feeds reais e autorização.",
    "Automação futura via GitHub Actions ou rotina local agendada, mas mantendo segredos fora do frontend.",
    "Possível migração futura para framework com build step, mas atualmente manter GitHub Pages simples.",
])

add_heading(doc, "9) Regras Importantes", 1)
add_bullets(doc, [
    ("Nunca expor segredos: ", "scripts/.env, API keys, feed URLs privados e ficheiros CSV/GZ temporários não vão para GitHub."),
    ("Publicar dados processados: ", "o frontend só deve receber data/*.js públicos, sem chaves."),
    ("Performance primeiro: ", "manter site leve; evitar carregar feeds ou dados enormes no browser."),
    ("Separação dados/interface: ", "scripts geram dados; index.html consome dados; não misturar segredos ou lógica de download no frontend."),
    ("Não inventar reviews: ", "para SEO, aggregateRating/review só devem existir se forem reais."),
    ("Validar antes de publicar: ", "correr report-quality.js e report-duplicate-candidates.js antes de enviar para GitHub."),
    ("Manter histórico de duplicados: ", "não apagar data/duplicate-decisions.json sem motivo; é a memória manual do sistema."),
    ("Preferir feed a scraping: ", "não fazer scraping quando o acordo/loja só autoriza feed afiliado."),
])

add_heading(doc, "10) Linguagem e Padrões de Código", 1)
add_para(doc, "O código é JavaScript pragmático, sem build tool obrigatório no frontend. O index.html contém HTML, CSS e React/Babel embutido. Os scripts Node.js usam CommonJS, dotenv e módulos utilitários locais. A prioridade é clareza operacional e capacidade de manter o site estático em GitHub Pages.")
add_bullets(doc, [
    ("Naming: ", "scripts generate-<loja>.js, data/<loja>-data.js e variáveis window.PADELCOST_<LOJA>_PRODUCTS."),
    ("Normalização: ", "nomes e categorias são normalizados por scripts/name-normalization.js, scripts/brand-normalization.js e regras por loja."),
    ("Categorias: ", "usar ids internos em minúsculas: raquetes, sapatilhas, sacos, acessorios, roupa."),
    ("Preços: ", "números em euros como Number; apresentação formatada no frontend com vírgula e símbolo €."),
    ("Lojas: ", "stores[] deve ser ordenável por price e STORE_PRIORITY."),
    ("Comentários: ", "usar comentários só onde a lógica precisa de contexto; evitar ruído."),
    ("Validação: ", "usar node --check para scripts alterados e report-quality.js para dados."),
])

add_heading(doc, "11) Runbook Operacional", 1)
add_heading(doc, "11.1 Preparação local", 2)
add_code_block(doc, [
    "cd \"/Users/ruibarba/Desktop/files 2/scripts\"",
    "npm install",
    "cp .env.example .env",
    "# editar .env com chaves privadas e IDs de feed",
])

add_heading(doc, "11.2 Rotina semanal sem Atmosfera Sport", 2)
add_code_block(doc, [
    "cd \"/Users/ruibarba/Desktop/files 2/scripts\"",
    "npm run update:weekly:no-atmosfera",
])
add_para(doc, "Esta rotina corre geradores por loja, merge, qualidade e SEO, mas não deve regenerar a base da Atmosfera enquanto o feed estiver em revisão.")

add_heading(doc, "11.3 Validação manual recomendada", 2)
add_code_block(doc, [
    "cd \"/Users/ruibarba/Desktop/files 2\"",
    "node --check scripts/merge-offers.js",
    "node --check scripts/generate-seo-pages.js",
    "node scripts/report-quality.js",
    "node scripts/report-duplicate-candidates.js",
    "node scripts/generate-duplicate-review-page.js",
])

add_heading(doc, "11.4 Ficheiros importantes para GitHub após integração Deporvillage", 2)
add_bullets(doc, [
    "index.html",
    "lojas.html",
    "sitemap.xml",
    "data/products-data.js",
    "data/deporvillage-data.js",
    "data/seo-pages-data.js",
    "data/duplicate-decisions.json, se houver novas decisões manuais",
    "logos/deporvillage-logo.png",
    "scripts/generate-deporvillage.js",
    "scripts/merge-offers.js",
    "scripts/generate-seo-pages.js",
    "scripts/update-weekly-no-atmosfera.js",
    "scripts/package.json",
    "scripts/package-lock.json, se npm install alterar dependências",
    "scripts/.env.example",
])
add_callout(doc, "Não publicar", "Não enviar scripts/.env, scripts/.env.*, node_modules, CSVs, GZs, logs, ficheiros temporários ou URLs privados com API key.")

add_heading(doc, "11.5 Google Search Console", 2)
add_bullets(doc, [
    "Sitemap atual é /sitemap.xml e já foi submetido com sucesso no Search Console.",
    "Depois de publicar sitemap.xml atualizado, o Google tende a reler sozinho; também se pode reenviar o mesmo sitemap se necessário.",
    "Avisos de aggregateRating/review em fragmentos de produto são melhorias não críticas, não erros bloqueantes.",
    "Não carregar em “validar correção” antes de publicar alterações que resolvam de facto o problema.",
    "Páginas “detetadas atualmente não indexadas” podem ser normais em sites novos; priorizar qualidade, sitemap, interlinking e conteúdo útil.",
])

add_heading(doc, "12) Resumo Essencial", 1)
add_bullets(doc, [
    "PadelCost é um comparador estático de preços de padel publicado em GitHub Pages.",
    "A pasta local de trabalho é /Users/ruibarba/Desktop/files 2.",
    "O repositório de publicação é PadelCost VR2 e o upload é manual.",
    "A app principal está em index.html e consome data/products-data.js.",
    "Não há backend; os feeds são processados localmente por scripts Node.js.",
    "As chaves privadas ficam em scripts/.env e nunca vão para GitHub.",
    "Lojas atuais: Padel Market, Zona de Padel, Padel Proshop PT, Atmosfera Sport, Forum Sport ES, Decathlon ES, Adidas Padel e Deporvillage.",
    "Categorias públicas: raquetes, sapatilhas, sacos e acessórios.",
    "O merge junta ofertas por EAN, GTIN, MPN, assinatura segura e decisões manuais.",
    "O catálogo atual tem 2544 produtos, 2353 públicos e 461 produtos com 2+ lojas.",
    "As páginas SEO são geradas em categoria/ e produto/ por scripts/generate-seo-pages.js.",
    "Há uma ferramenta local para rever duplicados: duplicados-revisao.html.",
    "Antes de publicar, correr relatório de qualidade, relatório de duplicados e regenerar SEO/sitemap.",
    "Atmosfera Sport deve continuar fora da rotina completa até o feed ser revisto.",
    "Próximo trabalho provável: rever duplicados, SEO e integrar novas lojas por feed autorizado.",
])

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("PadelCost — passaporte técnico interno")

doc.save(OUTPUT)
print(OUTPUT)
